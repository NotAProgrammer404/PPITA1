#!/usr/bin/env python3
"""
Layout-Preserving OCR using Google Gemini Vision API
Preserves document layout and handles images intelligently.
"""

import os
import sys
import argparse
import re
from pathlib import Path
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables
load_dotenv()

# Configure Gemini API
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    print("Error: GOOGLE_API_KEY not found in .env file")
    print("Please copy .env.example to .env and add your API key")
    sys.exit(1)

genai.configure(api_key=GOOGLE_API_KEY)


def process_image(image_path: str, output_path: str = None) -> str:
    """
    Process an image using Gemini Vision API to extract text with layout preservation.
    
    Args:
        image_path: Path to the image file
        output_path: Optional path for output file
        
    Returns:
        Extracted text in Markdown format
    """
    try:
        # Load and validate image
        img = Image.open(image_path)
        print(f"📄 Processing: {image_path}")
        print(f"   Image size: {img.size[0]}x{img.size[1]} pixels")
        
        # Initialize Gemini model with vision capabilities
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        # Craft prompt for layout preservation
        prompt = """Analyze this handwritten document/notes image and extract ALL text while preserving the EXACT layout, structure, and spatial relationships.

CRITICAL INSTRUCTIONS:

1. LAYOUT PRESERVATION:
   - Maintain the EXACT spatial layout (left/right columns, top/bottom sections)
   - Preserve boxes, borders, and enclosed content together
   - Keep related content grouped as it appears visually
   - Use horizontal rules (---) to separate distinct sections
   - Use tables or columns when content is side-by-side

2. HEADINGS & STRUCTURE:
   - Use # for main headings (usually in red or underlined)
   - Use ## for subheadings and section titles
   - Use ### for sub-sections
   - Preserve numbering systems exactly as shown

3. BOXES & BORDERS:
   - For content inside boxes/borders, use this format:
     ```
     ┌─────────────────────┐
     │ Content inside box  │
     └─────────────────────┘
     ```
   - Or use blockquotes for boxed content:
     > **[Boxed Content]**
     > Content here

4. DIAGRAMS & VISUAL ELEMENTS:
   - Describe diagrams in detail with their position
   - Use ASCII art for simple diagrams/arrows where possible
   - Format: ![Diagram: detailed description](diagram_placeholder.jpg)
   - Preserve arrows: → ← ↑ ↓ ⇒ ⇔
   - Show connections between concepts

5. FORMULAS & EQUATIONS:
   - Keep mathematical/chemical formulas on separate lines
   - Preserve subscripts (H₂O) and superscripts where possible
   - Box important formulas:
     ```
     Formula: x/m = KP^(1/n)
     ```

6. TEXT FORMATTING:
   - **Bold** for emphasized/highlighted text (often in red/colored)
   - *Italic* for definitions or important terms
   - `code blocks` for formulas, equations, or technical notation
   - Bullet points (•) for lists
   - Numbered lists (1., 2., 3.) where numbered

7. HANDWRITING SPECIFICS:
   - Extract ALL handwritten text accurately
   - Preserve abbreviations and shorthand
   - Keep symbols: ∝, ≈, ≠, ≤, ≥, Δ, α, β, etc.
   - Maintain original spelling/grammar

8. SPATIAL RELATIONSHIPS:
   - If content is in two columns, use a table:
     | Left Column | Right Column |
     |-------------|--------------|
     | Content     | Content      |
   - Preserve the reading flow as it appears visually

9. COMPLETE EXTRACTION:
   - Extract EVERY piece of text, including small notes
   - Include all examples (eg., e.g.)
   - Capture all labels, annotations, and side notes

Output ONLY the formatted content in a way that preserves the visual layout when rendered. No preamble or explanations."""

        # Generate content
        print("🤖 Analyzing with Gemini Vision...")
        response = model.generate_content([prompt, img])
        
        # Extract text
        extracted_text = response.text
        
        # Save to file if output path specified
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
            print(f"✅ Saved to: {output_path}")
        
        return extracted_text
        
    except FileNotFoundError:
        print(f"❌ Error: File not found: {image_path}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error processing image: {str(e)}")
        sys.exit(1)


def process_pdf(pdf_path: str, output_path: str = None) -> str:
    """
    Process a PDF file by converting pages to images and processing each.
    
    Args:
        pdf_path: Path to the PDF file
        output_path: Optional path for output file
        
    Returns:
        Extracted text from all pages in Markdown format
    """
    try:
        from pdf2image import convert_from_path
        
        print(f"📚 Processing PDF: {pdf_path}")
        
        # Convert PDF to images
        print("🔄 Converting PDF pages to images...")
        images = convert_from_path(pdf_path)
        print(f"   Found {len(images)} page(s)")
        
        # Initialize Gemini model
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        all_pages_text = []
        
        # Process each page
        for i, img in enumerate(images, 1):
            print(f"\n📄 Processing page {i}/{len(images)}...")
            
            prompt = """Analyze this handwritten document/notes image and extract ALL text while preserving the EXACT layout, structure, and spatial relationships.

CRITICAL INSTRUCTIONS:

1. LAYOUT PRESERVATION:
   - Maintain the EXACT spatial layout (left/right columns, top/bottom sections)
   - Preserve boxes, borders, and enclosed content together
   - Keep related content grouped as it appears visually
   - Use horizontal rules (---) to separate distinct sections
   - Use tables or columns when content is side-by-side

2. HEADINGS & STRUCTURE:
   - Use # for main headings (usually in red or underlined)
   - Use ## for subheadings and section titles
   - Use ### for sub-sections
   - Preserve numbering systems exactly as shown

3. BOXES & BORDERS:
   - For content inside boxes/borders, use this format:
     ```
     ┌─────────────────────┐
     │ Content inside box  │
     └─────────────────────┘
     ```
   - Or use blockquotes for boxed content:
     > **[Boxed Content]**
     > Content here

4. DIAGRAMS & VISUAL ELEMENTS:
   - Describe diagrams in detail with their position
   - Use ASCII art for simple diagrams/arrows where possible
   - Format: ![Diagram: detailed description](diagram_placeholder.jpg)
   - Preserve arrows: → ← ↑ ↓ ⇒ ⇔
   - Show connections between concepts

5. FORMULAS & EQUATIONS:
   - Keep mathematical/chemical formulas on separate lines
   - Preserve subscripts (H₂O) and superscripts where possible
   - Box important formulas:
     ```
     Formula: x/m = KP^(1/n)
     ```

6. TEXT FORMATTING:
   - **Bold** for emphasized/highlighted text (often in red/colored)
   - *Italic* for definitions or important terms
   - `code blocks` for formulas, equations, or technical notation
   - Bullet points (•) for lists
   - Numbered lists (1., 2., 3.) where numbered

7. HANDWRITING SPECIFICS:
   - Extract ALL handwritten text accurately
   - Preserve abbreviations and shorthand
   - Keep symbols: ∝, ≈, ≠, ≤, ≥, Δ, α, β, etc.
   - Maintain original spelling/grammar

8. SPATIAL RELATIONSHIPS:
   - If content is in two columns, use a table:
     | Left Column | Right Column |
     |-------------|--------------|
     | Content     | Content      |
   - Preserve the reading flow as it appears visually

9. COMPLETE EXTRACTION:
   - Extract EVERY piece of text, including small notes
   - Include all examples (eg., e.g.)
   - Capture all labels, annotations, and side notes

Output ONLY the formatted content in a way that preserves the visual layout when rendered. No preamble or explanations."""

            response = model.generate_content([prompt, img])
            page_text = response.text
            
            # Add page separator
            if i > 1:
                all_pages_text.append(f"\n\n---\n**Page {i}**\n---\n\n")
            else:
                all_pages_text.append(f"**Page {i}**\n\n")
            
            all_pages_text.append(page_text)
        
        # Combine all pages
        full_text = "".join(all_pages_text)
        
        # Save to file if output path specified
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            print(f"\n✅ Saved to: {output_path}")
        
        return full_text
        
    except ImportError:
        print("❌ Error: pdf2image not installed or poppler-utils missing")
        print("Install with: pip install pdf2image")
        print("On Linux, also run: sudo apt-get install poppler-utils")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error processing PDF: {str(e)}")
        sys.exit(1)


def markdown_to_docx(markdown_text: str, output_path: str):
    """
    Convert Markdown text to a Word document (.docx).
    
    Args:
        markdown_text: Markdown formatted text
        output_path: Path for the output .docx file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Handle bullet points
        elif line.strip().startswith('*   ') or line.strip().startswith('- '):
            text = line.strip()[4:] if line.strip().startswith('*   ') else line.strip()[2:]
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
        
        # Handle tables
        elif '|' in line and line.strip():
            if not in_table:
                in_table = True
                table_rows = []
            
            # Skip separator lines
            if re.match(r'^\s*\|[\s\-:]+\|', line):
                i += 1
                continue
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_rows.append(cells)
            
            # Check if next line is still part of table
            if i + 1 < len(lines) and '|' not in lines[i + 1]:
                # Create table
                if table_rows:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Bold header row
                            if row_idx == 0:
                                cell.paragraphs[0].runs[0].font.bold = True
                    
                    doc.add_paragraph()  # Add spacing after table
                in_table = False
                table_rows = []
        
        # Handle images (placeholders)
        elif line.strip().startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                alt_text = match.group(1)
                p = doc.add_paragraph()
                p.add_run(f'[Image: {alt_text}]').italic = True
        
        # Handle regular paragraphs
        elif line.strip() and not line.startswith('#'):
            # Process inline formatting
            text = line.strip()
            p = doc.add_paragraph()
            
            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Handle italic
                    italic_parts = re.split(r'(\*.*?\*)', part)
                    for ipart in italic_parts:
                        if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                            run = p.add_run(ipart[1:-1])
                            run.italic = True
                        else:
                            # Handle code
                            code_parts = re.split(r'(`.*?`)', ipart)
                            for cpart in code_parts:
                                if cpart.startswith('`') and cpart.endswith('`'):
                                    run = p.add_run(cpart[1:-1])
                                    run.font.name = 'Courier New'
                                elif cpart:
                                    p.add_run(cpart)
        
        # Empty lines
        else:
            if i > 0 and lines[i-1].strip():  # Only add paragraph if previous line had content
                doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"✅ Word document saved to: {output_path}")



def main():
    """Main entry point for the OCR tool."""
    parser = argparse.ArgumentParser(
        description='Layout-Preserving OCR using Gemini Vision API',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python ocr.py document.png
  python ocr.py report.pdf -o output.md
  python ocr.py image.jpg --output result.md
        """
    )
    
    parser.add_argument('input', help='Input image or PDF file')
    parser.add_argument('-o', '--output', help='Output Markdown file (optional)')
    
    args = parser.parse_args()
    
    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"❌ Error: Input file not found: {args.input}")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = args.output
    else:
        output_path = input_path.stem + '_ocr.md'
    
    # Process based on file type
    file_ext = input_path.suffix.lower()
    
    print("=" * 60)
    print("🔍 Layout-Preserving OCR with Gemini Vision")
    print("=" * 60)
    
    if file_ext == '.pdf':
        result = process_pdf(str(input_path), output_path)
    elif file_ext in ['.png', '.jpg', '.jpeg', '.webp', '.gif']:
        result = process_image(str(input_path), output_path)
    else:
        print(f"❌ Error: Unsupported file type: {file_ext}")
        print("Supported formats: PNG, JPG, JPEG, WEBP, GIF, PDF")
        sys.exit(1)
    
    print("\n" + "=" * 60)
    print("✨ Processing complete!")
    print("=" * 60)
    
    # Create Word document from Markdown
    docx_path = output_path.replace('.md', '.docx') if output_path.endswith('.md') else output_path + '.docx'
    print(f"\n📄 Creating Word document...")
    markdown_to_docx(result, docx_path)
    
    # Display preview
    print("\n📝 Preview (first 500 characters):")
    print("-" * 60)
    print(result[:500])
    if len(result) > 500:
        print("...")
    print("-" * 60)


if __name__ == '__main__':
    main()
