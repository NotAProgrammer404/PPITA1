#!/usr/bin/env python3
"""
Enhanced Layout-Preserving OCR using Google Gemini Vision API
Extracts images, provides JSON structure with positions, and embeds images in DOCX.
"""

import os
import sys
import argparse
import re
import json
from pathlib import Path
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables
load_dotenv()

# Configure Gemini API
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    print("Error: GOOGLE_API_KEY not found in .env file")
    print("Please copy .env.example to .env and add your API key")
    sys.exit(1)

genai.configure(api_key=GOOGLE_API_KEY)


def process_image_with_structure(image_path: str, extract_images: bool = True, images_dir: str = "images") -> dict:
    """
    Process an image using Gemini Vision API to extract structured content with positions.
    
    Args:
        image_path: Path to the image file
        extract_images: Whether to extract image regions
        images_dir: Directory to save extracted images
        
    Returns:
        Dictionary with structured document data
    """
    try:
        # Load and validate image
        img = Image.open(image_path)
        print(f"📄 Processing: {image_path}")
        print(f"   Image size: {img.size[0]}x{img.size[1]} pixels")
        
        # Create images directory if needed
        if extract_images:
            os.makedirs(images_dir, exist_ok=True)
        
        # Initialize Gemini model with vision capabilities
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        # Craft prompt for structured JSON output
        prompt = """Analyze this document/handwritten notes image and extract ALL content in a structured JSON format.

CRITICAL INSTRUCTIONS:

1. Return ONLY valid JSON (no markdown code blocks, no explanations)
2. Identify ALL elements: headings, paragraphs, images/diagrams, tables, formulas, lists
3. For each element, provide:
   - type: "heading", "paragraph", "image", "table", "formula", "list", "box"
   - content: the actual text or description
   - level: (for headings only) 1, 2, or 3
   - position: approximate relative position as "top", "middle", "bottom", "left", "right", "center"
   - bbox: approximate bounding box as {"x": 0-100, "y": 0-100, "width": 0-100, "height": 0-100} in percentage
   - is_math: true/false (true if content contains mathematical expressions, equations, formulas, or symbols)
   - style: formatting information including:
     * font_size: "small", "medium", "large", "xlarge" (relative size)
     * font_weight: "normal", "bold"
     * font_style: "normal", "italic"
     * text_color: "black", "red", "blue", "green", etc. (if colored)
     * underline: true/false
     * alignment: "left", "center", "right"

4. SPELL CHECKING & CLEANING:
   - Correct any obvious spelling errors in the extracted text (including HEADINGS)
   - Ensure the content is coherent and makes sense
   - For scientific/technical terms, preserve the correct spelling
   - CLEAN UP HEADINGS: Remove markdown symbols like '#', '*', or underlining characters from the 'content' field. Example: "# Adsorption" -> "Adsorption"
   - Do NOT add content that isn't in the image
   - Do NOT hallucinate or make up information

5. MATHEMATICAL CONTENT DETECTION:
   - If text contains equations, formulas, or mathematical symbols (=, +, -, ×, ÷, ∝, α, β, Δ, ∫, ∑, √, ², ³, subscripts, superscripts, fractions), mark as type: "formula" AND set is_math: true
   - Even if the math is bold or italic, still classify as "formula"
   - Examples of math: "x/m = KP^(1/n)", "ΔH > TAS", "H₂O", "Tc = 8a/27Rb"
   - Mathematical relationships like "α" (proportional to) should be marked as formulas

6. LAYOUT PRESERVATION:
   - Maintain the exact spatial layout of the original document
   - Preserve column structure if present
   - Keep relative positions accurate (top/middle/bottom, left/center/right)
   - If content is in two columns, mark positions accordingly
   - Preserve page boundaries - don't add extra content

7. For images/diagrams:
   - Provide detailed description
   - Indicate if it's a diagram, chart, illustration, or photo
   - Describe what it shows
   - ESTIMATE SIZE: Provide 'image_size': "small" (icon/symbol), "medium" (column width), "large" (full page width)
   - BOUNDING BOXES: Ensure image bounding boxes are GENEROUS and capture the FULL diagram/image including all labels, axes, and legends. Do not crop tightly.

8. Preserve reading order (top to bottom, left to right within columns)

JSON Structure:
{
  "document": {
    "title": "Main document title if present",
    "width": <original width in pixels>,
    "height": <original height in pixels>,
    "elements": [
      {
        "id": 1,
        "type": "heading",
        "level": 1,
        "content": "Heading text",
        "position": "top-center",
        "bbox": {"x": 10, "y": 5, "width": 80, "height": 8},
        "is_math": false,
        "style": {
          "font_size": "xlarge",
          "font_weight": "bold",
          "font_style": "normal",
          "text_color": "black",
          "underline": true,
          "alignment": "center"
        }
      },
      {
        "id": 2,
        "type": "image",
        "description": "Detailed description of the image/diagram",
        "image_type": "diagram",
        "image_size": "medium",
        "position": "middle-left",
        "bbox": {"x": 5, "y": 30, "width": 40, "height": 25}
      },
      {
        "id": 3,
        "type": "formula",
        "content": "x/m = KP^(1/n)",
        "position": "middle-center",
        "bbox": {"x": 40, "y": 35, "width": 20, "height": 5},
        "is_math": true,
        "style": {
          "font_size": "medium",
          "font_weight": "bold",
          "font_style": "normal",
          "text_color": "black",
          "underline": false,
          "alignment": "center"
        }
      },
      {
        "id": 4,
        "type": "paragraph",
        "content": "Text content here",
        "position": "middle-right",
        "bbox": {"x": 50, "y": 30, "width": 45, "height": 20},
        "is_math": false,
        "style": {
          "font_size": "medium",
          "font_weight": "normal",
          "font_style": "normal",
          "text_color": "black",
          "underline": false,
          "alignment": "left"
        }
      }
    ]
  }
}

Extract EVERYTHING and return ONLY the JSON."""

        # Generate content
        print("🤖 Analyzing with Gemini Vision...")
        response = model.generate_content([prompt, img])
        
        # Extract and parse JSON
        response_text = response.text.strip()
        
        # Remove markdown code blocks if present
        if response_text.startswith('```'):
            response_text = re.sub(r'^```(?:json)?\n', '', response_text)
            response_text = re.sub(r'\n```$', '', response_text)
        
        try:
            structured_data = json.loads(response_text)
        except json.JSONDecodeError as e:
            print(f"⚠️  Warning: Could not parse JSON response. Error: {e}")
            print("Response text:")
            print(response_text[:500])
            # Fallback: create basic structure
            structured_data = {
                "document": {
                    "title": "Extracted Document",
                    "width": img.size[0],
                    "height": img.size[1],
                    "elements": [
                        {
                            "id": 1,
                            "type": "paragraph",
                            "content": response_text,
                            "position": "full",
                            "bbox": {"x": 0, "y": 0, "width": 100, "height": 100}
                        }
                    ]
                }
            }
        
        # Add actual dimensions
        structured_data["document"]["width"] = img.size[0]
        structured_data["document"]["height"] = img.size[1]
        
        # Extract images if requested
        if extract_images:
            image_counter = 1
            for element in structured_data["document"]["elements"]:
                if element["type"] == "image" and "bbox" in element:
                    bbox = element["bbox"]
                    
                    # Convert percentage to pixels
                    x = int(bbox["x"] * img.size[0] / 100)
                    y = int(bbox["y"] * img.size[1] / 100)
                    w = int(bbox["width"] * img.size[0] / 100)
                    h = int(bbox["height"] * img.size[1] / 100)
                    
                    # Add padding to ensure complete capture (2% of image dimension)
                    pad_x = int(img.size[0] * 0.02)
                    pad_y = int(img.size[1] * 0.02)
                    
                    # Calculate new coordinates with padding, keeping within bounds
                    x_start = max(0, x - pad_x)
                    y_start = max(0, y - pad_y)
                    x_end = min(img.size[0], x + w + pad_x)
                    y_end = min(img.size[1], y + h + pad_y)
                    
                    # Crop image with padding
                    try:
                        cropped = img.crop((x_start, y_start, x_end, y_end))
                        image_filename = f"image_{image_counter:03d}.png"
                        image_path_full = os.path.join(images_dir, image_filename)
                        cropped.save(image_path_full)
                        
                        # Add image path to element
                        element["image_path"] = image_path_full
                        print(f"   ✂️  Extracted image {image_counter}: {image_filename} (padded)")
                        image_counter += 1
                    except Exception as e:
                        print(f"   ⚠️  Could not extract image {image_counter}: {e}")
        
        return structured_data
        
    except FileNotFoundError:
        print(f"❌ Error: File not found: {image_path}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error processing image: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


def json_to_docx(json_data: dict, output_path: str):
    """
    Convert JSON structured data to a Word document with embedded images and styling.
    Preserves layout to fit content on minimal pages.
    
    Args:
        json_data: Structured document data
        output_path: Path for the output .docx file
    """
    doc = Document()
    
    # Set very tight margins for layout preservation (single page)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
    
    # Set default font - very compact
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(8)  # Very small for single-page fit
    
    # Reduce spacing between paragraphs to minimum
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(1)  # Absolute minimum spacing
    paragraph_format.line_spacing = 0.9  # Tighter than single spacing
    
    # Font size mapping - all reduced
    font_size_map = {
        "small": Pt(7),
        "medium": Pt(8),
        "large": Pt(10),
        "xlarge": Pt(14)
    }
    
    # Text color mapping
    from docx.shared import RGBColor
    color_map = {
        "black": RGBColor(0, 0, 0),
        "red": RGBColor(255, 0, 0),
        "blue": RGBColor(0, 0, 255),
        "green": RGBColor(0, 128, 0),
        "orange": RGBColor(255, 165, 0),
        "purple": RGBColor(128, 0, 128)
    }
    
    # Alignment mapping
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }
    
    # Add title if present
    if "title" in json_data["document"] and json_data["document"]["title"]:
        title = doc.add_heading(json_data["document"]["title"], level=0)
        title.paragraph_format.space_after = Pt(6)  # Minimal spacing after title
    
    # Process elements in order
    for element in json_data["document"]["elements"]:
        element_type = element.get("type", "paragraph")
        style_info = element.get("style", {})
        
        if element_type == "heading":
            level = element.get("level", 1)
            content = element.get("content", "")
            heading = doc.add_heading(content, level=level)
            heading.paragraph_format.space_after = Pt(2)  # Minimal spacing
            
            # Apply styles to heading runs
            if heading.runs and style_info:
                for run in heading.runs:
                    if "font_size" in style_info:
                        run.font.size = font_size_map.get(style_info["font_size"], Pt(12))
                    if "font_weight" in style_info and style_info["font_weight"] == "bold":
                        run.bold = True
                    if "text_color" in style_info:
                        run.font.color.rgb = color_map.get(style_info["text_color"], RGBColor(0, 0, 0))
                    if "underline" in style_info and style_info["underline"]:
                        run.underline = True
                    if "font_style" in style_info and style_info["font_style"] == "italic":
                        run.italic = True
            
            # Apply alignment if specified
            if "alignment" in style_info:
                heading.alignment = alignment_map.get(style_info["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
        
        elif element_type == "paragraph":
            content = element.get("content", "")
            if content.strip():
                p = doc.add_paragraph()
                
                # Apply alignment
                if "alignment" in style_info:
                    p.alignment = alignment_map.get(style_info["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
                
                # Process inline formatting
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', content)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                    else:
                        run = p.add_run(part)
                    
                    # Apply style from JSON
                    if style_info:
                        if "font_size" in style_info:
                            run.font.size = font_size_map.get(style_info["font_size"], Pt(11))
                        if "font_weight" in style_info and style_info["font_weight"] == "bold":
                            run.bold = True
                        if "font_style" in style_info and style_info["font_style"] == "italic":
                            run.italic = True
                        if "text_color" in style_info:
                            run.font.color.rgb = color_map.get(style_info["text_color"], RGBColor(0, 0, 0))
                        if "underline" in style_info and style_info["underline"]:
                            run.underline = True
        
        elif element_type == "image":
            # Add image if path exists
            if "image_path" in element and os.path.exists(element["image_path"]):
                try:
                    # Dynamic image size based on prompt analysis
                    size_category = element.get("image_size", "medium")
                    
                    if size_category == "small":
                        width = Inches(1.0)
                    elif size_category == "large":
                        width = Inches(3.5)  # Full width for single page compact layout
                    else:  # medium
                        width = Inches(2.0)
                        
                    doc.add_picture(element["image_path"], width=width)
                    # Skip caption to save space
                except Exception as e:
                    print(f"⚠️  Could not embed image: {e}")
            # Skip text placeholders to save space
        
        elif element_type == "formula":
            content = element.get("content", "")
            is_math = element.get("is_math", True)  # Default to true for formulas
            p = doc.add_paragraph()
            run = p.add_run(content)
            
            # Use math-friendly font for formulas
            if is_math:
                run.font.name = 'Cambria Math'  # Better for mathematical symbols
            else:
                run.font.name = 'Courier New'
            
            run.font.size = Pt(11)
            
            # Apply style from JSON if available
            if style_info:
                if "font_size" in style_info:
                    run.font.size = font_size_map.get(style_info["font_size"], Pt(11))
                if "font_weight" in style_info and style_info["font_weight"] == "bold":
                    run.bold = True
                if "alignment" in style_info:
                    p.alignment = alignment_map.get(style_info["alignment"], WD_ALIGN_PARAGRAPH.CENTER)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif element_type == "list":
            items = element.get("content", [])
            for item in items:
                # Handle both string and dictionary list items
                item_text = ""
                if isinstance(item, dict):
                    item_text = item.get("item", item.get("content", str(item)))
                else:
                    item_text = str(item)
                    
                p = doc.add_paragraph(item_text, style='List Bullet')
                
                # Apply style if specified
                if style_info and p.runs:
                    if "font_size" in style_info:
                        p.runs[0].font.size = font_size_map.get(style_info["font_size"], Pt(11))
        
        elif element_type == "box":
            content = element.get("content", "")
            p = doc.add_paragraph()
            run = p.add_run(f"┌{'─' * 50}┐\n│ {content}\n└{'─' * 50}┘")
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        elif element_type == "table":
            # Handle table data if present
            rows = element.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        if row_idx == 0:
                            cell.paragraphs[0].runs[0].font.bold = True
                doc.add_paragraph()  # Spacing
    
    # Save document
    doc.save(output_path)
    print(f"✅ Word document saved to: {output_path}")


def main():
    """Main entry point for the enhanced OCR tool."""
    parser = argparse.ArgumentParser(
        description='Enhanced Layout-Preserving OCR with Image Extraction',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python ocr_enhanced.py document.png
  python ocr_enhanced.py notes.jpg --json-output structure.json
  python ocr_enhanced.py image.png --extract-images --images-dir extracted_images
        """
    )
    
    parser.add_argument('input', help='Input image file')
    parser.add_argument('-j', '--json-output', help='Output JSON file (optional)')
    parser.add_argument('-o', '--output', help='Output DOCX file (optional)')
    parser.add_argument('--extract-images', action='store_true', default=True,
                        help='Extract image regions (default: True)')
    parser.add_argument('--images-dir', default='images',
                        help='Directory to save extracted images (default: images)')
    
    args = parser.parse_args()
    
    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"❌ Error: Input file not found: {args.input}")
        sys.exit(1)
    
    # Determine output paths
    base_name = input_path.stem
    json_output = args.json_output or f"{base_name}_structure.json"
    docx_output = args.output or f"{base_name}_enhanced.docx"
    
    print("=" * 60)
    print("🔍 Enhanced OCR with Image Extraction")
    print("=" * 60)
    
    # Process image
    structured_data = process_image_with_structure(
        str(input_path),
        extract_images=args.extract_images,
        images_dir=args.images_dir
    )
    
    # Save JSON
    with open(json_output, 'w', encoding='utf-8') as f:
        json.dump(structured_data, f, indent=2, ensure_ascii=False)
    print(f"✅ JSON structure saved to: {json_output}")
    
    # Create DOCX
    print(f"\n📄 Creating enhanced Word document...")
    json_to_docx(structured_data, docx_output)
    
    print("\n" + "=" * 60)
    print("✨ Processing complete!")
    print("=" * 60)
    
    # Display summary
    num_elements = len(structured_data["document"]["elements"])
    num_images = sum(1 for e in structured_data["document"]["elements"] if e["type"] == "image")
    
    print(f"\n📊 Summary:")
    print(f"   Total elements: {num_elements}")
    print(f"   Images/diagrams: {num_images}")
    print(f"   JSON output: {json_output}")
    print(f"   DOCX output: {docx_output}")


if __name__ == '__main__':
    main()
